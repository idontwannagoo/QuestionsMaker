#!/usr/bin/env python
#-*- coding:utf-8 -*-


from tkinter import *
from tkinter.font import Font
from tkinter.ttk import *
from tkinter.messagebox import *
#import tkinter.filedialog as tkFileDialog
#import tkinter.simpledialog as tkSimpleDialog    #askstring()
import random
import json
import re
from retry import retry
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import collections

class Application_ui(Frame):
    #这个类仅实现界面生成功能，具体事件处理代码在子类Application中。
    def __init__(self, master=None):
        Frame.__init__(self, master)
        self.master.title('Form1')
        # To center the window on the screen.
        ws = self.master.winfo_screenwidth()
        hs = self.master.winfo_screenheight()
        x = (ws / 2) - (392 / 2)
        y = (hs / 2) - (431 / 2)
        self.master.geometry('%dx%d+%d+%d' % (392,431,x,y))
        self.createWidgets()

    def createWidgets(self):
        self.top = self.winfo_toplevel()

        self.style = Style()

        self.Check2Var = IntVar(value=0)
        self.style.configure('TCheck2.TCheckbutton', font=('微软雅黑',9,'bold'))
        self.Check2 = Checkbutton(self.top, text='题号', variable=self.Check2Var, style='TCheck2.TCheckbutton')
        self.Check2.place(relx=0.388, rely=0.705, relwidth=0.125, relheight=0.039)

        self.Check1Var = IntVar(value=0)
        self.style.configure('TCheck1.TCheckbutton', font=('微软雅黑',9,'bold'))
        self.Check1 = Checkbutton(self.top, text='打乱顺序', variable=self.Check1Var, style='TCheck1.TCheckbutton')
        self.Check1.place(relx=0.122, rely=0.705, relwidth=0.247, relheight=0.039)

        self.Command1Var = StringVar(value='出题')
        self.style.configure('TCommand1.TButton', font=('微软雅黑',9,'bold'))
        self.Command1 = Button(self.top, text='出题', textvariable=self.Command1Var, command=self.Command1_Cmd, style='TCommand1.TButton')
        self.Command1.place(relx=0.082, rely=0.798, relwidth=0.839, relheight=0.114)

        self.topRadioVar = StringVar()
        self.style.configure('TOption3.TRadiobutton', font=('微软雅黑',9,'bold'))
        self.Option3 = Radiobutton(self.top, text='随机挖空', value='Option3', variable=self.topRadioVar, style='TOption3.TRadiobutton')
        self.Option3.place(relx=0.653, rely=0.557, relwidth=0.247, relheight=0.046)

        self.style.configure('TOption2.TRadiobutton', font=('微软雅黑',9,'bold'))
        self.Option2 = Radiobutton(self.top, text='挖空下联', value='Option2', variable=self.topRadioVar, style='TOption2.TRadiobutton')
        self.Option2.place(relx=0.388, rely=0.557, relwidth=0.247, relheight=0.046)

        self.style.configure('TOption1.TRadiobutton', font=('微软雅黑',9,'bold'))
        self.Option1 = Radiobutton(self.top, text='挖空上联', value='Option1', variable=self.topRadioVar, style='TOption1.TRadiobutton')
        self.Option1.place(relx=0.122, rely=0.557, relwidth=0.247, relheight=0.046)

        self.Combo1List = ['七年级上册', '七年级下册', '八年级上册',
                           '八年级下册']
        self.Combo1Var = StringVar(value='请在此选择预置诗题')
        self.Combo1 = Combobox(self.top, textvariable=self.Combo1Var, values=self.Combo1List, font=('微软雅黑', 9, 'bold'))
        self.Combo1.bind('<<ComboboxSelected>>', self.Add_Listbox_Cmd)
        self.Combo1.place(relx=0.122, rely=0.093, relwidth=0.798)

        self.Label2Var = StringVar(value='请选择出题方式：')
        self.style.configure('TLabel2.TLabel', anchor='w', font=('微软雅黑',9,'bold'))
        self.Label2 = Label(self.top, textvariable=self.Label2Var, style='TLabel2.TLabel')
        self.Label2.place(relx=0.082, rely=0.483, relwidth=0.245, relheight=0.039)

        self.Lable1Var = StringVar(value='请选择年段等级：')
        self.style.configure('TLable1.TLabel', anchor='w', font=('微软雅黑',9,'bold'))
        self.Lable1 = Label(self.top, textvariable=self.Lable1Var, style='TLable1.TLabel')
        self.Lable1.place(relx=0.082, rely=0.037, relwidth=0.245, relheight=0.039)

        self.Label3Var = StringVar(value='附加选项：')
        self.style.configure('TLabel1.TLabel', anchor='w', font=('微软雅黑',9,'bold'))
        self.Label1 = Label(self.top, textvariable=self.Label3Var, style='TLabel1.TLabel')
        self.Label1.place(relx=0.082, rely=0.631, relwidth=0.153, relheight=0.039)

        self.Check3Var = IntVar(value=0)
        self.style.configure('TCheck3.TCheckbutton', font=('微软雅黑',9,'bold'))
        self.Check3 = Checkbutton(self.top, text='附加作者和诗题', variable=self.Check3Var, style='TCheck3.TCheckbutton')
        self.Check3.place(relx=0.653, rely=0.705, relwidth=0.268, relheight=0.046)

        self.PoemsList = []
        self.PoemsListBoxFont = Font(font=('微软雅黑', 9))
        self.PoemsListBox = Listbox(self.top, font=self.PoemsListBoxFont, selectmode=MULTIPLE)
        # for each in self.PoemsList:
        #     self.PoemsListBox.insert(END, each)
        self.PoemsListBox.place(relx=0.122, rely=0.167, relwidth=0.798, relheight=0.288)
class Application(Application_ui):
    #这个类实现具体的事件处理回调函数。界面生成代码在Application_ui中。
    def __init__(self, master=None):
        Application_ui.__init__(self, master)
        with open('data.json', 'r', encoding='UTF-8') as f:
            self.data = json.load(f)
        self.topRadioVar.set("Option3")
        self.Check1Var.set(1)
        self.Check2Var.set(1)
        self.Check3Var.set(1)
    def Command1_Cmd(self, event=None):
        self.DataProcessing()
        self.make_questions()
        self.shuffle()
        self.add_number()
        self.write_docx()

    def remove_empty(self, list):
        for each in list:
            if each == '':
                list.remove(each)
        return list
    def DataProcessing(self):
        poems = self.data[self.Combo1Var.get()]['诗文'].split('\n')
        self.listboxDisplayList = []
        self.wordtitleDict = collections.OrderedDict()
        wordTampList = []
        titleTamp = ''
        for poem in poems:
            attr = poem.split('|')
            words = attr[0]
            title = attr[2] + attr[3]
            page = attr[4]

            if title != titleTamp:
                wordTampList.clear()
            titleTamp = title
            wordTampList.append(words)
            self.wordtitleDict[title] = wordTampList[:]
            self.listboxDisplayList.append(title)
        self.title_list = list(self.wordtitleDict.keys())
    def Add_Listbox_Cmd(self, event=None):
        self.DataProcessing()
        self.PoemsListBox.delete(0, END)  # 清空列表显示
        listboxDisplayOdr = list(set(self.listboxDisplayList))
        listboxDisplayOdr.sort(key=self.listboxDisplayList.index)
        for each in listboxDisplayOdr:
            self.PoemsListBox.insert(END, each)


    def make_questions(self):
        '''数据的读取和解析'''
        listboxIdxTuple = self.PoemsListBox.curselection()
        wordtitleDictCopy = collections.OrderedDict()
        for each in listboxIdxTuple:
            value = self.wordtitleDict.pop(self.title_list[each])
            wordtitleDictCopy[self.title_list[each]] = value
        self.questions = []
        self.answers = []
        for title, words in wordtitleDictCopy.items():
            # 制作挖完空的题目和未挖空的答案列表
            for word in words:
                word_list = re.split(r"([。！，？；])", word)
                word_list_no_pun = re.split(r"[。！，？；]", word)
                # 删除空元素
                word_list = self.remove_empty(word_list)
                word_list_no_pun = self.remove_empty(word_list_no_pun)
                # 生成答案
                answer = ''.join(word_list)
                self.add_title(list=self.answers, str=answer, title=title)
                # 判断选项设置挖空位置
                if self.topRadioVar.get() == 'Option1':
                    item_index = 0
                elif self.topRadioVar.get() == 'Option2':
                    item_index = 1
                else:
                    item_index = random.randint(0, len(word_list_no_pun) - 1)
                # 执行挖空
                word_list[word_list.index(word_list_no_pun[item_index])] = len(word_list_no_pun[item_index]) * '_' * 3
                question = ''.join(word_list)  # 列表组织成字符串（诗文）
                # 字符串加入列表
                self.add_title(list=self.questions, str=question, title=title)
    def add_title(self, list, str, title):
        if self.Check3Var.get() == 1:
            list.append(str + '（' + title + '）')
        else:
            list.append(str)

    def shuffle(self):
        '''打乱诗题'''
        if self.Check1Var.get() == 1:
            # 相同顺序打乱
            k = list(zip(self.questions, self.answers))
            random.shuffle(k)
            self.questions[:], self.answers[:] = zip(*k)

    def add_number(self):
        '''在诗题前加上序号'''
        if self.Check2Var.get() == 1:
            for num in range(len(self.questions)):
                self.questions[num] = str(num+1) + '. ' + self.questions[num]
                self.answers[num] = str(num+1) + '. ' + self.answers[num]

    def create_docx(self, list):
        '''实现文档的创建和内容写入，不包括保存文档'''
        # 创建文档
        doc = Document()
        # 创建段落
        title = doc.add_paragraph()  # 标题
        psg = doc.add_paragraph()  # 正文
        # 标题内容设置
        title_run = title.add_run(self.Combo1Var.get() + '古诗练习')
        # 正文内容设置
        for each in list:
            psg.add_run(each + '\n')
        # 标题样式设计
        title_run.bold = True  # 粗体
        title_run.font.size = Pt(16)  # 字号
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
        # 字体设置（东亚字体）
        title_run.font.name = u'黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        # 正文
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        return doc  # 返回结果

    def write_docx(self):
        question_doc = self.create_docx(self.questions)
        question_doc.save(self.Combo1Var.get() + '诗题.docx')
        answers_doc = self.create_docx(self.answers)
        answers_doc.save(self.Combo1Var.get() + '诗题答案.docx')
        showinfo('完成', '文件已经输出到软件根目录，分为答案和题目两个文件。')


if __name__ == "__main__":
    top = Tk()
    top.resizable(width=False, height=False)
    Application(top).mainloop()
